# -*- coding: utf-8 -*-
"""把工作台提示词 .md 转换成 .docx（Word），保留标题/列表/加粗/代码块结构。
用法：python md2docx.py <输入.md> <输出.docx>
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CODE_FONT = "Consolas"

def add_inline(paragraph, text):
    """解析 **加粗** 与 `代码` 并写入同一段落（混合 run）。"""
    # 用占位拆分：先按 ** 和 ` 切，但简单起见用正则交替匹配
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)

def convert(md_path, docx_path):
    doc = Document()
    # 基础中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_buf = []

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        s = line.strip()
        if s == "":
            continue
        if s == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BBBBBB")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue
        if s.startswith("# "):
            add_inline(doc.add_heading(s[2:].strip(), level=1), "")
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=4)
            continue
        if re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, s[2:].strip())
            continue
        if re.match(r"^\d+[.、]\s+", s):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+[.、]\s+", "", s))
            continue
        p = doc.add_paragraph()
        add_inline(p, s)

    flush_code()
    doc.save(docx_path)
    print(f"  已生成: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python md2docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
